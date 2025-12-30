#!/usr/bin/env python3
"""
Generate a Word document with screenshots and captions.
Scales images to fit A4 page width.
"""

import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Paths
SCRIPT_DIR = Path(__file__).parent
PROJECT_DIR = SCRIPT_DIR.parent
SCREENSHOT_DIR = PROJECT_DIR / 'screenshots'
OUTPUT_FILE = PROJECT_DIR / 'ISP-Portal-Screenshots.docx'

# A4 page dimensions (with margins)
# A4 is 21cm x 29.7cm, with 2.54cm margins on each side = ~16cm usable width
A4_WIDTH = Cm(16)

def create_document():
    """Create the Word document with screenshots."""

    # Load manifest
    manifest_path = SCREENSHOT_DIR / 'manifest.json'
    if not manifest_path.exists():
        print(f"Error: Manifest not found at {manifest_path}")
        print("Run take-screenshots.js first.")
        return

    with open(manifest_path) as f:
        manifest = json.load(f)

    screenshots = manifest.get('screenshots', [])

    if not screenshots:
        print("No screenshots found in manifest.")
        return

    # Create document
    doc = Document()

    # Set up styles
    title_style = doc.styles['Title']
    title_style.font.size = Pt(24)
    title_style.font.bold = True

    # Add title page
    title = doc.add_heading('BTRC ISP Self-Care Portal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('User Interface Screenshots')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)

    doc.add_paragraph()

    date_para = doc.add_paragraph('December 2024')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Add table of contents placeholder
    toc_heading = doc.add_heading('Table of Contents', level=1)

    # Track unique page titles for TOC
    seen_titles = set()
    toc_entries = []

    for item in screenshots:
        base_title = item['title'].split(' - ')[0].split(' (continued')[0]
        if base_title not in seen_titles:
            seen_titles.add(base_title)
            toc_entries.append(base_title)

    for i, entry in enumerate(toc_entries, 1):
        doc.add_paragraph(f"{i}. {entry}")

    doc.add_page_break()

    # Add screenshots with captions
    current_section = None
    figure_num = 1

    for item in screenshots:
        filename = item['filename']
        title = item['title']

        # Determine section from title
        section_name = title.split(' - ')[0].split(' (continued')[0]

        # Add section heading if new section
        if section_name != current_section:
            if current_section is not None:
                doc.add_page_break()
            doc.add_heading(section_name, level=1)
            current_section = section_name

        # Add the image
        image_path = SCREENSHOT_DIR / filename
        if not image_path.exists():
            print(f"Warning: Image not found: {image_path}")
            continue

        # Add image scaled to A4 width
        try:
            doc.add_picture(str(image_path), width=A4_WIDTH)

            # Center the image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add caption
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run(f"Figure {figure_num}: {title}")
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True

            doc.add_paragraph()  # Add spacing

            print(f"Added: Figure {figure_num} - {title}")
            figure_num += 1

        except Exception as e:
            print(f"Error adding image {filename}: {e}")

    # Save document
    doc.save(OUTPUT_FILE)
    print(f"\n✓ Document saved to: {OUTPUT_FILE}")
    print(f"✓ Total figures: {figure_num - 1}")

if __name__ == '__main__':
    create_document()
