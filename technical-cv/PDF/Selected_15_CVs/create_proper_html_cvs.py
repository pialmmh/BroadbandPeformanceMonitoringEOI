import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
import fitz
import os
import re
from html import escape

# Position assignments based on careful matching (19 people)
positions = {
    'Md. Ariful Haque': 'Solution Architect / System Analyst',
    'Mohammad Jahirul Islam': 'Network & QoS Expert',
    'Mahbubul Islam Shaikat': 'Data Integration / Interoperability Expert',
    'Mohammad Arif Iftekhar': 'Senior Software Engineer',
    'Mahathir': 'Senior Software Engineer',
    'Md. Abdullah Al-Asif': 'Mobile Application Developer',
    'Ashik': 'Mobile Application Developer',
    'Abdullah Al Ferdous': 'Database Administrator / Data Engineer',
    'Gazi Md Tawsif': 'Security Expert',
    'MD IBRAHIM KHALILULLAH': 'Quality Assurance Engineer',
    'Syed Mohammad Easin': 'UI/UX Designer',
    'Nargish': 'Business Analyst',
    'Md. Anarul Islam': 'Trainer / Support Engineer',
    'Rakib Hasan Sazib': 'Trainer / Support Engineer',
    'S. M. Asif Zawad': 'Trainer / Support Engineer',
    'Apple': 'Senior Software Engineer',
    'tanvir': 'Software Developer',
    'Mahfuz Alam': 'Support Engineer',
    'Md Mahmudur Rahman': 'Support Engineer'
}

# Name mapping from filename
name_mapping = {
    'Md. Ariful Haque': 'Md. Ariful Haque',
    'Mohammad Jahirul Islam': 'Mohammad Jahirul Islam',
    'Mahbubul Islam Shaikat': 'Mahbubul Islam Shaikat',
    'Mohammad Arif Iftekhar': 'Mohammad Arif Iftekhar',
    'Mahathir CV': 'Mahathir Mohammad Bishal',
    'Mahathir': 'Mahathir Mohammad Bishal',
    'Md. Abdullah Al-Asif': 'Md. Abdullah Al-Asif',
    'Ashik CV': 'Ashik Mahmud',
    'Ashik': 'Ashik Mahmud',
    'Abdullah Al Ferdous': 'Abdullah Al Ferdous',
    'Gazi Md Tawsif': 'Gazi Md Tawsif',
    'MD IBRAHIM KHALILULLAH MUNNA': 'MD Ibrahim Khalilullah (Munna)',
    'MD IBRAHIM KHALILULLAH': 'MD Ibrahim Khalilullah (Munna)',
    'Syed Mohammad Easin CV': 'Syed Mohammad Easin',
    'Syed Mohammad Easin': 'Syed Mohammad Easin',
    'Nargish CV': 'Nargish Sultana',
    'Nargish': 'Nargish Sultana',
    'Md. Anarul Islam': 'Md. Anarul Islam',
    'Rakib Hasan Sazib': 'Rakib Hasan Sazib',
    'S. M. Asif Zawad': 'S. M. Asif Zawad',
    'MD._Apple_Mahmud': 'Md. Apple Mahmud',
    'Apple': 'Md. Apple Mahmud',
    'tanvirCV': 'Tanvir Hasan',
    'tanvir': 'Tanvir Hasan',
    'Mahfuz Alam': 'Mahfuz Alam',
    'Md Mahmudur Rahman': 'Md Mahmudur Rahman'
}

def get_position(filename):
    for name, pos in positions.items():
        if name.lower() in filename.lower():
            return pos
    return 'Technical Expert'

def get_proper_name(filename):
    """Get proper name from filename"""
    base_name = filename.replace('.pdf', '').replace('.docx', '').strip()
    for key, value in name_mapping.items():
        if key.lower() in base_name.lower():
            return value
    name = base_name.replace(' CV', '').replace('_CV', '').replace(' cv', '')
    name = re.sub(r'_+', ' ', name)
    return name.strip()

def extract_docx_structured(filepath):
    """Extract data from DOCX files that already follow the template format"""
    doc = Document(filepath)
    data = {
        'name': '',
        'dob': '',
        'nationality': 'Bangladeshi',
        'membership': 'N/A',
        'education': '',
        'training': '',
        'languages': '',
        'work_experience': '',
        'employment_records': [],
        'specializations': '',
        'projects': '',
        'skills': ''
    }

    if len(doc.tables) == 0:
        return data

    table = doc.tables[0]
    current_section = None  # 'employment', 'specializations', 'projects', 'skills'

    for row_idx, row in enumerate(table.rows):
        # Get unique cell contents
        cells = []
        seen_texts = set()
        for cell in row.cells:
            text = cell.text.strip()
            if text and text not in seen_texts:
                seen_texts.add(text)
                cells.append(text)

        if not cells:
            continue

        first_cell = cells[0].lower()
        full_row_text = ' '.join(cells)
        first_cell_full = cells[0]

        # Detect section changes - check for employment first (various formats)
        if '11.' in first_cell and 'employment' in first_cell:
            current_section = 'employment'
            continue
        elif 'employment record' in first_cell:
            current_section = 'employment'
            continue
        elif '11.' in first_cell or 'subject matter' in first_cell:
            current_section = 'specializations'
            text = full_row_text
            if '11.' in text:
                parts = text.split('11.', 1)
                if len(parts) > 1:
                    text = parts[1].strip()
                    if text.lower().startswith('subject matter'):
                        text = text[len('subject matter'):].strip()
                        if text.lower().startswith('specializations'):
                            text = text[len('specializations'):].strip()
            data['specializations'] = text
            continue
        elif '12.' in first_cell or 'work undertaken' in first_cell or 'additional experience' in first_cell:
            current_section = 'projects'
            text = full_row_text
            if '12.' in text:
                parts = text.split('12.', 1)
                if len(parts) > 1:
                    text = parts[1].strip()
            data['projects'] = text
            continue
        elif '13.' in first_cell or 'computer skill' in first_cell:
            current_section = 'skills'
            text = full_row_text
            if '13.' in text:
                parts = text.split('13.', 1)
                if len(parts) > 1:
                    text = parts[1].strip()
                    if text.lower().startswith('computer skill'):
                        text = text[len('computer skills'):].strip()
            data['skills'] = text
            continue

        # Add to current section if in one
        if current_section == 'specializations':
            data['specializations'] += '\n' + full_row_text
            continue
        elif current_section == 'projects':
            data['projects'] += '\n' + full_row_text
            continue
        elif current_section == 'skills':
            data['skills'] += '\n' + full_row_text
            continue
        elif current_section == 'employment':
            # Employment entry rows - parse "Employer : From To" format
            if 'employer' in first_cell and 'from' in full_row_text.lower():
                continue  # Skip header row
            elif 'experience:' in first_cell.lower():
                # Experience description
                if data['employment_records']:
                    data['employment_records'][-1]['experience'] = full_row_text
            elif len(cells) >= 1:
                # Check if this is a company/position row
                # Format: "Company Name\n    Position" or just company with dates
                company_text = first_cell_full
                dates = ''
                for cell in cells[1:]:
                    if cell != ':':
                        dates += ' ' + cell
                dates = dates.strip()

                # Check for "Till Date" or "Present" indicating current job
                if company_text and ('Ltd' in company_text or 'Limited' in company_text or
                                     'Inc' in company_text or 'Company' in company_text or
                                     '\n' in first_cell_full):
                    data['employment_records'].append({
                        'employer': company_text,
                        'dates': dates
                    })
            continue

        # Extract basic fields (before employment section)
        if 'name of staff' in first_cell:
            for cell in cells[1:]:
                if cell != ':' and cell:
                    data['name'] = cell
                    break
        elif 'date of birth' in first_cell:
            for cell in cells[1:]:
                if cell != ':' and cell:
                    data['dob'] = cell
                    break
        elif 'nationality' in first_cell:
            for cell in cells[1:]:
                if cell != ':' and cell:
                    data['nationality'] = cell
                    break
        elif 'membership' in first_cell:
            for cell in cells[1:]:
                if cell != ':' and cell:
                    data['membership'] = cell
                    break
        elif '6.' in first_cell and 'education' in first_cell:
            for cell in cells[1:]:
                if cell != ':' and cell:
                    data['education'] = cell
                    break
        elif '7.' in first_cell and ('training' in first_cell or 'skill' in first_cell):
            for cell in cells[1:]:
                if cell != ':' and cell:
                    data['skills'] = cell
                    break
        elif 'languages' in first_cell and 'proficiency' in first_cell:
            for cell in cells[1:]:
                if cell != ':' and cell:
                    data['languages'] = cell
                    break
        elif 'countries of work' in first_cell:
            for cell in cells[1:]:
                if cell != ':' and cell:
                    data['work_experience'] = cell
                    break
        elif '10.' in first_cell and ('training' in first_cell or 'certification' in first_cell):
            # Training/Certifications section
            for cell in cells[1:]:
                if cell != ':' and cell:
                    data['training'] = cell
                    break
        elif '10.' in first_cell and 'employment' in first_cell:
            current_section = 'employment'
        elif current_section is None and re.match(r'^\d+\.', first_cell_full.strip()):
            # Numbered employment entry outside of detected section
            # Skip if it's a field label like "Proposed Position", "Name of Staff", etc.
            if any(skip in first_cell.lower() for skip in ['proposed position', 'name of staff', 'date of birth', 'nationality', 'membership', 'education', 'training', 'language', 'countries']):
                continue
            emp_entry = first_cell_full
            dates = ''
            for cell in cells[1:]:
                if cell != ':':
                    dates += ' ' + cell
            data['employment_records'].append({
                'employer': emp_entry,
                'dates': dates.strip()
            })

    return data

def extract_pdf_content(filepath, filename):
    """Extract data from PDF files and map to template format"""
    doc = fitz.open(filepath)
    full_text = ''

    for page in doc:
        full_text += page.get_text() + '\n'

    doc.close()

    data = {
        'name': get_proper_name(filename),
        'dob': '',
        'nationality': 'Bangladeshi',
        'membership': 'N/A',
        'education': '',
        'training': '',
        'languages': 'Language\tSpeaking\tReading\tWriting\nEnglish\tExcellent\tExcellent\tExcellent\nBangla\tExcellent\tExcellent\tExcellent',
        'work_experience': 'Bangladesh',
        'employment_records': [],
        'specializations': '',
        'projects': '',
        'skills': ''
    }

    lines = full_text.split('\n')

    # Parse sections - detect section headers more accurately
    current_section = None
    education_content = []
    training_content = []
    employment_content = []
    project_content = []
    skill_content = []
    summary_content = []
    certificate_content = []

    # Section header patterns (exact matches for infographic-style CVs)
    section_headers = {
        'EDUCATION': 'education',
        'WORK EXPERIENCE': 'employment',
        'PROFESSIONAL SKILL': 'skills',
        'PROFILE SUMMARY': 'summary',
        'CERTIFICATES': 'training',
        'CERTIFICATE': 'training',
        'CERTIFICATES & ONLINE PROGRAMME': 'training',
        'COURSES': 'courses',
        'REFERENCE': 'reference',
        'LANGUAGES': None  # Skip languages section header
    }

    for i, line in enumerate(lines):
        line_stripped = line.strip()
        line_lower = line_stripped.lower()

        if not line_stripped:
            continue

        # Check for exact section headers (common in infographic CVs)
        if line_stripped in section_headers:
            current_section = section_headers[line_stripped]
            continue

        # Detect DOB
        if 'date of birth' in line_lower:
            if ':' in line_stripped:
                data['dob'] = line_stripped.split(':', 1)[1].strip()
            elif i + 1 < len(lines) and lines[i + 1].strip():
                data['dob'] = lines[i + 1].strip()

        # Detect years of experience
        match = re.search(r'(\d+)\+?\s*years?\s*(of)?\s*(experience|it experience)?', line_lower)
        if match:
            data['work_experience'] = f"Bangladesh - {match.group(0)}"

        # Section detection for traditional formats
        if current_section is None:
            if 'education qualification' in line_lower or (line_lower == 'education'):
                current_section = 'education'
                continue
            elif 'professional certificate' in line_lower or ('certification' in line_lower and len(line_stripped) < 30):
                current_section = 'training'
                continue
            elif 'career summary' in line_lower or 'profile summary' in line_lower or 'objective' in line_lower:
                current_section = 'summary'
                continue
            elif 'completed project' in line_lower or 'project:' in line_lower:
                current_section = 'projects'
                continue
            elif 'working experience' in line_lower or ('work experience' in line_lower):
                current_section = 'employment'
                continue
            elif 'technical skill' in line_lower or 'professional skill' in line_lower or (line_lower == 'skill' or line_lower == 'skills'):
                current_section = 'skills'
                continue

        # Stop collecting at reference section
        if current_section == 'reference':
            continue

        # Collect content
        if current_section == 'education':
            education_content.append(line_stripped)
        elif current_section == 'training':
            training_content.append(line_stripped)
        elif current_section == 'courses':
            training_content.append(line_stripped)  # Add courses to training
        elif current_section == 'employment':
            employment_content.append(line_stripped)
        elif current_section == 'projects':
            project_content.append(line_stripped)
        elif current_section == 'skills':
            skill_content.append(line_stripped)
        elif current_section == 'summary':
            summary_content.append(line_stripped)

        # Also detect certification content that may appear outside of detected sections
        # (common in multi-column PDF layouts)
        cert_keywords = ['Red Hat', 'SUSE', 'Veeam', 'Cybersecurity', 'IBM -', 'Google Analytics',
                         'RH124', 'RH134', 'RH294', 'RH354', 'DO180', 'DO280', 'DO288', 'DO316']
        if any(kw in line_stripped for kw in cert_keywords):
            if line_stripped not in training_content:
                training_content.append(line_stripped)

    # Process education content - parse institution names with degrees
    # Also look for education items that may be mixed into other sections
    all_education_lines = list(education_content)

    # Check employment content for education items too (common in infographic CVs)
    for line in employment_content:
        if 'University' in line or 'College' in line or 'School' in line:
            all_education_lines.append(line)
        if 'Bachelor' in line or 'Masters' in line or 'B.Sc' in line or 'M.Sc' in line:
            all_education_lines.append(line)
        if 'Secondary School Certificate' in line or 'Higher Secondary' in line:
            all_education_lines.append(line)

    if all_education_lines:
        edu_parsed = []
        degrees = []
        institutions = []

        for line in all_education_lines:
            # Skip dates on their own lines, email addresses, phone numbers
            if re.match(r'^\d{4}$', line) or '@' in line or line.startswith('+'):
                continue
            if len(line) < 5:
                continue

            # Categorize lines
            if 'Bachelor' in line or 'Masters' in line or 'Certificate' in line or 'B.Sc' in line or 'M.Sc' in line:
                degrees.append(line)
            elif 'University' in line or 'College' in line or 'School' in line:
                institutions.append(line)
            else:
                edu_parsed.append(line)

        # Combine degrees and institutions - remove duplicates while preserving order
        combined_edu = []
        seen = set()
        for degree in degrees:
            if degree not in seen:
                combined_edu.append(degree)
                seen.add(degree)
        for inst in institutions:
            if inst not in seen:
                # Check if already included in a degree line
                if not any(inst in d for d in degrees):
                    combined_edu.append(inst)
                    seen.add(inst)
        for item in edu_parsed:
            if item not in seen:
                combined_edu.append(item)
                seen.add(item)

        data['education'] = '\n'.join(combined_edu[:15]) if combined_edu else ''

    # Process training/certificates
    if training_content:
        train_parsed = []
        for l in training_content:
            if len(l) < 5 or '@' in l:
                continue
            # Skip company names that might be in the section
            if 'Ltd' in l or 'Limited' in l or 'Company' in l or 'Pvt' in l:
                continue
            # Skip reference section content
            if 'General Manager' in l or 'VP of' in l or 'Chairman' in l:
                continue
            # Skip language entries
            if l in ['English (Fluent)', 'Bengali (Native)', 'English', 'Bengali']:
                continue
            train_parsed.append(l)
        data['training'] = '\n'.join(train_parsed[:30])

    # Process skills - stop at reference info, employment info, URLs
    if skill_content:
        skill_parsed = []
        for line in skill_content:
            # Stop at reference section markers
            if '@' in line or 'Dr.' in line or 'Chairman' in line or 'General Manager' in line:
                break
            # Skip employment info that got mixed in
            if '|' in line:  # "Position | Company" format
                continue
            if 'PRESENT' in line.upper():
                continue
            if re.match(r'^(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)\s+\d{4}', line, re.IGNORECASE):
                continue
            if line.startswith('http') or 'linkedin.com' in line.lower():
                continue
            if len(line) > 5:
                skill_parsed.append(line)
        data['skills'] = '\n'.join(skill_parsed[:20])

    data['projects'] = '\n'.join(project_content[:40]) if project_content else ''

    # Process summary - filter out contact info (name, email, phone, address)
    if summary_content:
        summary_parsed = []
        # Get the person's name to filter it out
        person_name = data.get('name', '').upper()
        name_parts = person_name.split() if person_name else []

        for line in summary_content:
            # Skip contact info
            if '@' in line or line.startswith('+'):
                continue
            # Skip if looks like an address (contains Road, Sector, Dhaka, etc.)
            if 'Road' in line or 'Sector' in line or 'Dhaka' in line or 'Uttara' in line:
                continue
            # Skip job titles that might be in header
            if 'Engineer' in line and len(line.split()) <= 4:
                continue
            # Skip if line is the person's name
            line_upper = line.upper()
            if name_parts and all(part in line_upper for part in name_parts[:2]):
                continue
            if len(line) > 10:  # Only substantial content
                summary_parsed.append(line)
        data['specializations'] = '\n'.join(summary_parsed[:20]) if summary_parsed else ''
    else:
        data['specializations'] = ''

    # Parse employment - look for "Position | Company" patterns
    # Also check skill_content for jobs that got mixed in
    all_employment_content = list(employment_content)
    for line in skill_content:
        if '|' in line or 'PRESENT' in line.upper():
            all_employment_content.append(line)

    if all_employment_content:
        emp_entries = []

        # For infographic-style CVs, extract only "Position | Company" lines
        job_lines = []
        date_lines = []

        for line in all_employment_content:
            # Skip contact info, education items
            if '@' in line or line.startswith('+') or 'Dhaka-' in line:
                continue
            if 'School' in line or 'College' in line or 'University' in line:
                continue
            if 'Certificate' in line or 'Bachelor' in line or 'Masters' in line or 'Secondary' in line:
                continue

            # Detect job entry pattern: "Title | Company"
            if '|' in line:
                job_lines.append(line)
            # Detect date patterns - with PRESENT or year range
            elif re.match(r'^(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)\s+\d{4}', line, re.IGNORECASE):
                date_lines.append(line)
            elif re.match(r'^\d{4}\s*-\s*(PRESENT|\d{4})', line, re.IGNORECASE):
                date_lines.append(line)
            elif 'PRESENT' in line.upper():
                date_lines.append(line)

        # Match job lines with dates - prioritize PRESENT dates
        present_date = None
        other_dates = []
        for d in date_lines:
            if 'PRESENT' in d.upper():
                present_date = d
            else:
                other_dates.append(d)

        # Create entries from job lines - assign PRESENT to most recent job
        for i, job in enumerate(job_lines):
            entry = job
            if i == 0 and present_date:
                # First job gets PRESENT date (assumed to be current)
                # But we need to check if this is actually the current job
                pass
            emp_entries.append(entry)

        # If we have a PRESENT date, find which job it belongs to
        # Look for the job that appears CLOSEST to (right before) PRESENT text
        if present_date:
            full_text = '\n'.join(all_employment_content)
            present_idx = full_text.upper().find('PRESENT')
            if present_idx > 0:
                # Find the job line that appears closest before PRESENT
                closest_job = None
                closest_distance = float('inf')
                for job in job_lines:
                    job_idx = full_text.find(job)
                    if job_idx >= 0 and job_idx < present_idx:
                        distance = present_idx - job_idx
                        if distance < closest_distance:
                            closest_distance = distance
                            closest_job = job

                if closest_job:
                    # Update the entry for this job
                    for i, entry in enumerate(emp_entries):
                        if entry == closest_job:
                            emp_entries[i] = closest_job + ' (Current - PRESENT)'
                            break

        # If no structured entries found, use raw content (filtered)
        if not emp_entries and employment_content:
            filtered = []
            for l in employment_content:
                if '@' in l or l.startswith('+') or 'Dhaka-' in l:
                    continue
                if 'School' in l or 'College' in l or 'University' in l:
                    continue
                if 'Certificate' in l or 'Bachelor' in l or 'Masters' in l:
                    continue
                if len(l) > 10:
                    filtered.append(l)
            if filtered:
                emp_entries = ['\n'.join(filtered[:15])]

        data['employment_records'] = emp_entries

    # If no skills found, use summary as specializations
    if not data['specializations'] and data['skills']:
        data['specializations'] = data['skills'][:500]

    return data

def get_current_employer(data):
    """Extract current employer from employment records"""
    if data.get('employment_records'):
        # First, try to find the record with "PRESENT" (current job)
        for record in data['employment_records']:
            record_text = str(record) if not isinstance(record, dict) else record.get('employer', '') + ' ' + record.get('dates', '')
            if 'PRESENT' in record_text.upper() or 'TILL DATE' in record_text.upper() or 'TILL NOW' in record_text.upper():
                # Found current job - extract company name
                if isinstance(record, dict):
                    employer = record.get('employer', '')
                    lines = employer.split('\n')
                    if lines:
                        company = lines[0].strip()
                        company = re.sub(r'^\d+\.\s*', '', company)
                        return company
                else:
                    lines = record_text.split('\n')
                    for line in lines:
                        if '|' in line:
                            parts = line.split('|')
                            if len(parts) >= 2:
                                company = parts[1].strip()
                                # Remove date/status suffixes
                                company = re.sub(r'\s*\(Current\s*-\s*PRESENT\)', '', company, flags=re.IGNORECASE)
                                company = re.sub(r'\s*\d{4}\s*-\s*(PRESENT|\d{4}).*', '', company, flags=re.IGNORECASE)
                                company = re.sub(r'\s*(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)\s+\d{4}.*', '', company, flags=re.IGNORECASE)
                                if company and len(company) > 3:
                                    return company.strip()

        # Fallback: use first record
        first_record = data['employment_records'][0]
        if isinstance(first_record, dict):
            employer = first_record.get('employer', '')
            lines = employer.split('\n')
            if lines:
                company = lines[0].strip()
                company = re.sub(r'^\d+\.\s*', '', company)
                return company
        else:
            record_text = str(first_record)
            lines = record_text.split('\n')
            if lines:
                first_line = lines[0].strip()

                # Handle "Position | Company" format
                if '|' in first_line:
                    parts = first_line.split('|')
                    if len(parts) >= 2:
                        company = parts[1].strip()
                        company = re.sub(r'\s*\d{4}\s*-\s*(PRESENT|\d{4}).*', '', company, flags=re.IGNORECASE)
                        if company:
                            return company

                # Try other patterns
                company = first_line
                company = re.sub(r'^\d+\.\s*', '', company)

                # Try to get just company name
                if 'Organization name:' in company:
                    company = company.split('Organization name:')[1].strip()

                if company and len(company) > 3:
                    return company

    return "CompTech Network System (Pvt) Ltd."

def generate_html(data, position):
    """Generate HTML following the sample template format"""

    # Get current employer for certification
    current_employer = get_current_employer(data)

    # Format employment records (simplified - no From/To columns)
    employment_html = ''
    if data.get('employment_records'):
        employment_html = '<table class="employment-table">\n'
        employment_html += '<tr><th>Employer / Position</th></tr>\n'

        for record in data['employment_records']:
            if isinstance(record, dict):
                emp = escape(record.get('employer', ''))
                dates = escape(record.get('dates', ''))
                exp = escape(record.get('experience', ''))

                # Combine employer and dates
                emp_with_dates = emp
                if dates:
                    emp_with_dates += f'<br><span style="font-size:10pt;color:#555;">Period: {dates}</span>'

                employment_html += f'<tr><td style="padding:10px;">{emp_with_dates}</td></tr>\n'
                if exp:
                    employment_html += f'<tr><td style="padding:5px 10px 10px 25px;font-style:italic;border-top:none;">{exp}</td></tr>\n'
            else:
                escaped_record = escape(str(record)).replace('\n', '<br>')
                employment_html += f'<tr><td style="padding:10px;">{escaped_record}</td></tr>\n'

        employment_html += '</table>'
    else:
        employment_html = ''

    # Build optional sections (only show if content exists)
    sections_html = ''

    # Section 10: Employment Records
    if employment_html:
        sections_html += f'''
    <div class="section-title">10. Employment Records:</div>
    <div class="section-content">
        {employment_html}
    </div>'''

    # Section 11: Subject Matter Specializations
    specializations = data.get('specializations', '').strip()
    if specializations:
        # Format specializations with proper HTML
        spec_escaped = escape(specializations)
        spec_lines = spec_escaped.split('\n')
        formatted_specs = ''
        for line in spec_lines:
            line = line.strip()
            if line:
                # Check if it's a category header (contains colon at end or specific keywords)
                if ':' in line and len(line.split(':')[0]) < 40:
                    parts = line.split(':', 1)
                    formatted_specs += f'<p style="margin:8px 0;"><strong>{parts[0]}:</strong> {parts[1] if len(parts) > 1 else ""}</p>\n'
                else:
                    formatted_specs += f'<p style="margin:5px 0 5px 15px;">{line}</p>\n'

        sections_html += f'''
    <div class="section-title">11. Subject Matter Specializations</div>
    <div class="section-content">
        {formatted_specs}
    </div>'''

    # Section 12: Work Undertaken / Projects - as a table
    projects = data.get('projects', '').strip()
    if projects:
        projects_escaped = escape(projects)

        # Build project table
        projects_table = '<table class="projects-table">\n'
        projects_table += '<tr><th style="width:30%">Project / Role</th><th style="width:70%">Description</th></tr>\n'

        # Try to parse project entries - look for clear numbered patterns like "(1)", "1.", "1)"
        project_entries = []

        # Pattern 1: Numbered entries like "(1) Project Name" or "1. Project Name"
        numbered_pattern = r'(?:^|\n)\s*(?:\((\d+)\)|(\d+)[\.\)])\s*'
        numbered_matches = list(re.finditer(numbered_pattern, projects_escaped))

        if len(numbered_matches) >= 2:
            # Found multiple numbered entries - use them as split points
            for i, match in enumerate(numbered_matches):
                start = match.end()
                end = numbered_matches[i + 1].start() if i + 1 < len(numbered_matches) else len(projects_escaped)
                content = projects_escaped[start:end].strip()
                if content:
                    # First line is project name, rest is description
                    lines = [l.strip() for l in content.split('\n') if l.strip()]
                    if lines:
                        project_name = lines[0]
                        description = ' '.join(lines[1:]) if len(lines) > 1 else ''
                        project_entries.append((project_name, description))

        # Pattern 2: Look for "Name of project:" markers
        if not project_entries and 'Name of project' in projects_escaped:
            parts = re.split(r'Name of project\s*:?\s*', projects_escaped, flags=re.IGNORECASE)
            for part in parts[1:]:  # Skip first empty part
                part = part.strip()
                if part:
                    lines = [l.strip() for l in part.split('\n') if l.strip()]
                    if lines:
                        project_name = lines[0]
                        description = ' '.join(lines[1:]) if len(lines) > 1 else ''
                        project_entries.append((project_name, description))

        # Pattern 3: Look for Role/Project patterns (like "Frontend Web Developer", "Web Developer at Upwork")
        if not project_entries:
            lines = [l.strip() for l in projects_escaped.split('\n') if l.strip()]

            # Filter out table header lines
            header_patterns = ['Project Name', 'Company Name', 'Project Role', 'Activities', 'Name of the assignment']
            filtered_lines = []
            for line in lines:
                # Skip if line is just headers
                is_header = False
                line_check = line.lower()
                for header in header_patterns:
                    if header.lower() in line_check and len(line) < 60:
                        # Check if line is mostly just headers
                        remaining = line_check
                        for h in header_patterns:
                            remaining = remaining.replace(h.lower(), '')
                        if len(remaining.strip()) < 10:
                            is_header = True
                            break
                if not is_header:
                    filtered_lines.append(line)

            lines = filtered_lines

            # Check if content has clear project/role structure
            role_keywords = ['Developer', 'Engineer', 'Designer', 'Analyst', 'Manager', 'Lead', 'Expert', 'Specialist', 'Architect']
            project_keywords = ['System', 'Application', 'Website', 'Portal', 'Platform', 'Solution', 'Infrastructure']

            current_entry = None
            current_desc = []

            for line in lines:
                # Check if this line looks like a title/role
                is_role_title = (
                    len(line) < 100 and
                    (any(kw in line for kw in role_keywords) or any(kw in line for kw in project_keywords)) and
                    not line.endswith(',') and
                    len(line.split()) <= 10
                )

                # Or if it's a clear short identifier (like project name - but not fragments)
                is_short_id = (
                    len(line) < 40 and
                    len(line.split()) <= 4 and
                    len(line.split()) >= 2 and  # At least 2 words
                    not line.endswith(',') and
                    not line.endswith('.') and
                    len(line) > 5  # Not too short
                )

                if is_role_title or (is_short_id and not current_entry):
                    # Save previous entry if exists
                    if current_entry and (current_entry or current_desc):
                        project_entries.append((current_entry, ' '.join(current_desc)))
                    current_entry = line
                    current_desc = []
                else:
                    current_desc.append(line)

            # Add last entry
            if current_entry:
                project_entries.append((current_entry, ' '.join(current_desc)))
            elif current_desc:
                # No structure found - join all content as description
                all_content = ' '.join(current_desc)
                project_entries.append(('Project Experience', all_content))

        # If still no entries, show content as single block
        if not project_entries:
            all_content = projects_escaped.replace('\n', ' ')
            all_content = re.sub(r'\s+', ' ', all_content).strip()
            project_entries.append(('Project Experience', all_content))

        # Post-process: check for fragmented entries and consolidate
        # If we have entries with very short/fragmented project names, consolidate them
        fragmented_count = sum(1 for p, d in project_entries if len(p) < 20 or p.endswith('and') or len(p.split()) <= 2)

        if fragmented_count > len(project_entries) / 2 and len(project_entries) > 2:
            # Too many fragmented entries - consolidate into a single description
            all_content = []
            for project_name, description in project_entries:
                all_content.append(project_name)
                if description:
                    all_content.append(description)
            combined = ' '.join(all_content)
            combined = re.sub(r'\s+', ' ', combined).strip()
            project_entries = [('Project Experience', combined)]

        # Generate table rows
        for project_name, description in project_entries:
            projects_table += f'<tr><td style="vertical-align:top;padding:8px;"><strong>{project_name}</strong></td>'
            projects_table += f'<td style="padding:8px;">{description}</td></tr>\n'

        projects_table += '</table>'

        sections_html += f'''
    <div class="section-title">12. Work Undertaken that Best Illustrates Capability to Handle the Tasks Assigned</div>
    <div class="section-content">
        {projects_table}
    </div>'''

    # Section 13: Computer Skills
    skills = data.get('skills', '').strip()
    if skills:
        # Format skills with proper HTML
        skills_escaped = escape(skills)
        skills_lines = skills_escaped.split('\n')
        formatted_skills = ''
        for line in skills_lines:
            line = line.strip()
            if line:
                if ':' in line and len(line.split(':')[0]) < 40:
                    parts = line.split(':', 1)
                    formatted_skills += f'<p style="margin:8px 0;"><strong>{parts[0]}:</strong> {parts[1] if len(parts) > 1 else ""}</p>\n'
                else:
                    formatted_skills += f'<p style="margin:5px 0;">{line}</p>\n'

        sections_html += f'''
    <div class="section-title">13. Computer Skills</div>
    <div class="section-content">
        {formatted_skills}
    </div>'''

    # Build main table rows (only show non-empty fields)
    main_rows = []

    # Required fields (always show)
    main_rows.append(f'''        <tr>
            <td class="label">1. Proposed Position</td>
            <td class="colon">:</td>
            <td class="value"><strong>{escape(position)}</strong></td>
        </tr>''')

    main_rows.append(f'''        <tr>
            <td class="label">2. Name of Staff</td>
            <td class="colon">:</td>
            <td class="value"><strong>{escape(data.get('name', ''))}</strong></td>
        </tr>''')

    # Optional fields (show only if content exists)
    if data.get('dob', '').strip():
        main_rows.append(f'''        <tr>
            <td class="label">3. Date of Birth</td>
            <td class="colon">:</td>
            <td class="value">{escape(data.get('dob', ''))}</td>
        </tr>''')

    main_rows.append(f'''        <tr>
            <td class="label">4. Nationality</td>
            <td class="colon">:</td>
            <td class="value">{escape(data.get('nationality', 'Bangladeshi'))}</td>
        </tr>''')

    if data.get('membership', '').strip() and data.get('membership', '').strip() != 'N/A':
        main_rows.append(f'''        <tr>
            <td class="label">5. Membership in Professional Society</td>
            <td class="colon">:</td>
            <td class="value">{escape(data.get('membership', 'N/A'))}</td>
        </tr>''')

    if data.get('education', '').strip():
        main_rows.append(f'''        <tr>
            <td class="label">6. Education</td>
            <td class="colon">:</td>
            <td class="value"><pre>{escape(data.get('education', ''))}</pre></td>
        </tr>''')

    if data.get('training', '').strip():
        main_rows.append(f'''        <tr>
            <td class="label">7. Other Training</td>
            <td class="colon">:</td>
            <td class="value"><pre>{escape(data.get('training', ''))}</pre></td>
        </tr>''')

    if data.get('languages', '').strip():
        # Format languages as a proper table
        lang_text = data.get('languages', '')
        lang_table = '<table class="lang-table"><tr><th>Language</th><th>Speaking</th><th>Reading</th><th>Writing</th></tr>'

        # Parse language data
        lines = lang_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line or 'Language' in line and 'Speaking' in line:
                continue  # Skip header line
            # Try to parse tab-separated or space-separated
            parts = line.replace('\t', '  ').split('  ')
            parts = [p.strip() for p in parts if p.strip()]
            if len(parts) >= 4:
                lang_table += f'<tr><td>{escape(parts[0])}</td><td>{escape(parts[1])}</td><td>{escape(parts[2])}</td><td>{escape(parts[3])}</td></tr>'
            elif len(parts) >= 1:
                # Just language name, assume Excellent for all
                lang_table += f'<tr><td>{escape(parts[0])}</td><td>Excellent</td><td>Excellent</td><td>Excellent</td></tr>'

        lang_table += '</table>'

        main_rows.append(f'''        <tr>
            <td class="label">8. Languages &amp; Degree of Proficiency</td>
            <td class="colon">:</td>
            <td class="value">{lang_table}</td>
        </tr>''')

    if data.get('work_experience', '').strip():
        main_rows.append(f'''        <tr>
            <td class="label">9. Countries of Work Experience</td>
            <td class="colon">:</td>
            <td class="value">{escape(data.get('work_experience', 'Bangladesh'))}</td>
        </tr>''')

    main_table_rows = '\n'.join(main_rows)

    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CV - {escape(data.get('name', 'Unknown'))}</title>
    <style>
        body {{
            font-family: 'Times New Roman', Times, serif;
            margin: 40px;
            line-height: 1.5;
            font-size: 11pt;
        }}
        .header {{
            text-align: center;
            margin-bottom: 20px;
        }}
        .main-table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 15px;
        }}
        .main-table td {{
            border: 1px solid #000;
            padding: 8px;
            vertical-align: top;
        }}
        .label {{
            width: 32%;
            font-weight: bold;
            background-color: #f5f5f5;
        }}
        .colon {{
            width: 3%;
            text-align: center;
            background-color: #f5f5f5;
        }}
        .value {{
            width: 65%;
        }}
        .section-title {{
            background-color: #f0f0f0;
            font-weight: bold;
            padding: 8px;
            margin-top: 15px;
            border: 1px solid #000;
            border-bottom: none;
        }}
        .section-content {{
            border: 1px solid #000;
            padding: 12px;
        }}
        .employment-table {{
            width: 100%;
            border-collapse: collapse;
        }}
        .employment-table th, .employment-table td {{
            border: 1px solid #ccc;
            padding: 6px;
            text-align: left;
            vertical-align: top;
        }}
        .employment-table th {{
            background-color: #f0f0f0;
            font-size: 10pt;
        }}
        .projects-table {{
            width: 100%;
            border-collapse: collapse;
        }}
        .projects-table th, .projects-table td {{
            border: 1px solid #ccc;
            padding: 8px;
            text-align: left;
            vertical-align: top;
        }}
        .projects-table th {{
            background-color: #f0f0f0;
            font-size: 10pt;
        }}
        .lang-table {{
            width: 100%;
            border-collapse: collapse;
        }}
        .lang-table th, .lang-table td {{
            border: 1px solid #ccc;
            padding: 5px 10px;
            text-align: center;
        }}
        .lang-table th {{
            background-color: #f5f5f5;
        }}
        .certification {{
            margin-top: 25px;
            padding: 15px;
            border: 1px solid #000;
        }}
        .signature-section {{
            margin-top: 20px;
        }}
        .checkbox-row {{
            margin: 15px 0;
        }}
        pre {{
            white-space: pre-wrap;
            font-family: 'Times New Roman', Times, serif;
            margin: 0;
            font-size: 11pt;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h2 style="margin:0;">CURRICULUM VITAE</h2>
    </div>

    <table class="main-table">
{main_table_rows}
    </table>
{sections_html}

    <div class="certification">
        <p><strong>Certification:</strong></p>
        <p>I, the undersigned, certify that (i) I was not a former employee of the client immediately before the submission of this proposal, and, (ii) I have not offered my CV to be proposed by a Firm other than this Consultant for this assignment, and (iii) to the best of my knowledge and belief, this bio data correctly describes myself, my qualifications, and my experience. I understand that any willful misstatement described herein may lead to my disqualification or dismissal, if engaged.</p>
        <p>I have been employed by <strong>{escape(current_employer)}</strong> as a consultant. Indicate Yes or No in the boxes below:</p>

        <div class="checkbox-row">
            <span style="margin-right: 30px;">YES &#9744;</span>
            <span>NO &#9744;</span>
        </div>

        <div class="signature-section">
            <p>Signature: _________________________</p>
            <p>Date of Signing: _____ / _____ / _______ (Day / Month / Year)</p>
        </div>
    </div>
</body>
</html>'''

    return html

def main():
    # Remove old HTML files
    if os.path.exists('HTML_CVs'):
        for f in os.listdir('HTML_CVs'):
            os.remove(os.path.join('HTML_CVs', f))

    os.makedirs('HTML_CVs', exist_ok=True)

    files = [f for f in os.listdir('.') if f.endswith(('.pdf', '.docx')) and 'Mustafa' not in f and 'create_' not in f]

    print(f'Processing {len(files)} CV files...\n')

    for filename in files:
        try:
            position = get_position(filename)
            proper_name = get_proper_name(filename)

            print(f'Processing: {filename}')
            print(f'  Name: {proper_name}')
            print(f'  Position: {position}')

            if filename.endswith('.docx'):
                data = extract_docx_structured(filename)
                if not data.get('name'):
                    data['name'] = proper_name
            else:
                data = extract_pdf_content(filename, filename)

            html_content = generate_html(data, position)

            safe_name = re.sub(r'[<>:"/\\|?*]', '', proper_name)
            output_file = f'HTML_CVs/{safe_name}.html'

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_content)

            print(f'  Created: {safe_name}.html\n')

        except Exception as e:
            print(f'  Error: {str(e)}')
            import traceback
            traceback.print_exc()
            print()

    print('\n=== Created HTML Files ===')
    for f in sorted(os.listdir('HTML_CVs')):
        print(f'  {f}')

if __name__ == '__main__':
    main()
