import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
import fitz
import os
import re
from html import escape

# Position assignments
positions = {
    'Md. Ariful Haque': 'Solution Architect / System Analyst',
    'Mohammad Jahirul Islam': 'Network & QoS Expert',
    'Mahbubul Islam Shaikat': 'Data Integration / Interoperability Expert',
    'Mohammad Arif Iftekhar': 'Senior Software Engineer',
    'Mahathir': 'Senior Software Engineer',
    'Md. Abdullah Al-Asif': 'Mobile Application Developer',
    'Ashik': 'Mobile Application Developer',
    'Abdullah Al Ferdous': 'Database Administrator / Data Engineer',
    'Gazi Md Tawsif': 'Security Expert',
    'MD IBRAHIM KHALILULLAH': 'Quality Assurance Engineer',
    'Syed Mohammad Easin': 'UI/UX Designer',
    'Nargish': 'Business Analyst',
    'Md. Anarul Islam': 'Trainer / Support Engineer',
    'Rakib Hasan Sazib': 'Trainer / Support Engineer',
    'S. M. Asif Zawad': 'Trainer / Support Engineer'
}

def get_position(filename):
    for name, pos in positions.items():
        if name.lower() in filename.lower():
            return pos
    return 'Technical Expert'

def get_name_from_filename(filename):
    name = filename.replace('.pdf', '').replace('.docx', '').replace(' CV', '').replace('_CV', '').replace(' cv', '')
    name = re.sub(r'_+', ' ', name)
    return name.strip()

def extract_docx_content(filepath):
    doc = Document(filepath)
    data = {}
    all_text = []

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = ' '.join(cells).lower()

            if 'name of staff' in row_text:
                for cell in cells:
                    if cell and 'name' not in cell.lower() and cell != ':':
                        data['name'] = cell
                        break
            elif 'date of birth' in row_text:
                for cell in cells:
                    if cell and 'date' not in cell.lower() and cell != ':':
                        data['dob'] = cell
                        break
            elif 'nationality' in row_text:
                for cell in cells:
                    if cell and 'nationality' not in cell.lower() and cell != ':':
                        data['nationality'] = cell
                        break
            elif 'education' in row_text and '6.' in cells[0]:
                for cell in cells:
                    if cell and 'education' not in cell.lower() and cell != ':':
                        data['education'] = cell
                        break
            elif 'countries of work' in row_text:
                for cell in cells:
                    if cell and 'countries' not in cell.lower() and cell != ':':
                        data['experience'] = cell
                        break
            elif 'language' in row_text and 'proficiency' in row_text:
                for cell in cells:
                    if cell and 'language' not in cell.lower() and cell != ':':
                        data['languages'] = cell
                        break
            elif 'training' in row_text and '7.' in cells[0]:
                for cell in cells:
                    if cell and 'training' not in cell.lower() and cell != ':':
                        data['training'] = cell
                        break
            elif 'membership' in row_text:
                for cell in cells:
                    if cell and 'membership' not in cell.lower() and cell != ':':
                        data['membership'] = cell
                        break

            # Collect all text
            for cell in row.cells:
                if cell.text.strip():
                    all_text.append(cell.text.strip())

    # Also collect paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text.strip())

    data['full_text'] = '\n'.join(all_text)
    return data

def extract_pdf_content(filepath):
    doc = fitz.open(filepath)
    data = {}
    full_text = ''

    for page in doc:
        full_text += page.get_text() + '\n'

    doc.close()
    data['full_text'] = full_text

    # Try to extract structured data from text
    lines = full_text.split('\n')
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if 'date of birth' in line_lower:
            if ':' in line:
                data['dob'] = line.split(':', 1)[1].strip()
            elif i + 1 < len(lines):
                data['dob'] = lines[i + 1].strip()
        elif 'nationality' in line_lower:
            if ':' in line:
                data['nationality'] = line.split(':', 1)[1].strip()

    return data

# HTML template
html_template = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CV - {name}</title>
    <style>
        body {{ font-family: 'Times New Roman', Times, serif; margin: 40px; line-height: 1.6; }}
        .header {{ text-align: center; margin-bottom: 20px; }}
        table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
        td, th {{ border: 1px solid #000; padding: 8px; text-align: left; vertical-align: top; }}
        .label {{ width: 30%; font-weight: bold; background-color: #f5f5f5; }}
        .colon {{ width: 5%; text-align: center; }}
        .value {{ width: 65%; }}
        .section-header {{ background-color: #e0e0e0; font-weight: bold; }}
        .certification {{ margin-top: 30px; padding: 20px; border: 1px solid #000; }}
        .signature {{ margin-top: 40px; }}
        h2 {{ color: #333; border-bottom: 2px solid #333; padding-bottom: 5px; }}
        pre {{ white-space: pre-wrap; font-family: 'Times New Roman', Times, serif; margin: 10px 0; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>CURRICULUM VITAE</h1>
    </div>

    <table>
        <tr>
            <td class="label">1. Proposed Position</td>
            <td class="colon">:</td>
            <td class="value"><strong>{position}</strong></td>
        </tr>
        <tr>
            <td class="label">2. Name of Staff</td>
            <td class="colon">:</td>
            <td class="value">{name}</td>
        </tr>
        <tr>
            <td class="label">3. Date of Birth</td>
            <td class="colon">:</td>
            <td class="value">{dob}</td>
        </tr>
        <tr>
            <td class="label">4. Nationality</td>
            <td class="colon">:</td>
            <td class="value">{nationality}</td>
        </tr>
        <tr>
            <td class="label">5. Membership in Professional Society</td>
            <td class="colon">:</td>
            <td class="value">{membership}</td>
        </tr>
        <tr>
            <td class="label">6. Education</td>
            <td class="colon">:</td>
            <td class="value">{education}</td>
        </tr>
        <tr>
            <td class="label">7. Other Training</td>
            <td class="colon">:</td>
            <td class="value">{training}</td>
        </tr>
        <tr>
            <td class="label">8. Languages &amp; Degree of Proficiency</td>
            <td class="colon">:</td>
            <td class="value">{languages}</td>
        </tr>
        <tr>
            <td class="label">9. Countries of Work Experience</td>
            <td class="colon">:</td>
            <td class="value">{work_countries}</td>
        </tr>
    </table>

    <h2>10. Employment Records</h2>
    <pre>{employment}</pre>

    <h2>11. Professional Experience &amp; Skills</h2>
    <pre>{skills}</pre>

    <div class="certification">
        <h3>Certification:</h3>
        <p>I, the undersigned, certify that (i) I was not a former employee of the client immediately before the submission of this proposal, and, (ii) I have not offered my CV to be proposed by a Firm other than this Consultant for this assignment, and (iii) to the best of my knowledge and belief, this bio data correctly describes myself, my qualifications, and my experience. I understand that any willful misstatement described herein may lead to my disqualification or dismissal, if engaged.</p>
        <p>I have been employed by <strong>Telcobright Limited</strong> as a consultant.</p>
    </div>

    <div class="signature">
        <table style="width: 50%; border: none;">
            <tr>
                <td style="border: none;">Signature: _____________________</td>
            </tr>
            <tr>
                <td style="border: none;">Date: _____________________</td>
            </tr>
        </table>
    </div>
</body>
</html>'''

def main():
    # Create output directory
    os.makedirs('HTML_CVs', exist_ok=True)

    # Process files
    files = [f for f in os.listdir('.') if f.endswith(('.pdf', '.docx')) and 'Mustafa' not in f]
    processed = 0

    for filename in files:
        try:
            name = get_name_from_filename(filename)
            position = get_position(filename)

            if filename.endswith('.docx'):
                data = extract_docx_content(filename)
            else:
                data = extract_pdf_content(filename)

            full_text = data.get('full_text', '')

            # Default values
            cv_data = {
                'name': data.get('name', name),
                'position': position,
                'dob': data.get('dob', ''),
                'nationality': data.get('nationality', 'Bangladeshi'),
                'membership': data.get('membership', 'N/A'),
                'education': data.get('education', ''),
                'training': data.get('training', ''),
                'languages': data.get('languages', 'English: Excellent, Bangla: Native'),
                'work_countries': data.get('experience', 'Bangladesh'),
                'employment': '',
                'skills': ''
            }

            # Parse full text for sections
            lines = full_text.split('\n')
            current_section = None
            employment_lines = []
            skill_lines = []
            education_lines = []
            training_lines = []

            for line in lines:
                line_stripped = line.strip()
                line_lower = line_stripped.lower()

                # Detect sections
                if any(x in line_lower for x in ['employment record', 'work experience', 'professional experience']):
                    current_section = 'employment'
                    continue
                elif any(x in line_lower for x in ['skill', 'technical expertise', 'computer skill', 'subject matter']):
                    current_section = 'skills'
                    continue
                elif 'education' in line_lower and len(line_stripped) < 50:
                    current_section = 'education'
                    continue
                elif any(x in line_lower for x in ['training', 'certification', 'certificate']) and len(line_stripped) < 50:
                    current_section = 'training'
                    continue

                if line_stripped:
                    if current_section == 'employment':
                        employment_lines.append(line_stripped)
                    elif current_section == 'skills':
                        skill_lines.append(line_stripped)
                    elif current_section == 'education':
                        education_lines.append(line_stripped)
                    elif current_section == 'training':
                        training_lines.append(line_stripped)

            if employment_lines:
                cv_data['employment'] = '\n'.join(employment_lines[:50])
            if skill_lines:
                cv_data['skills'] = '\n'.join(skill_lines[:30])
            if education_lines and not cv_data['education']:
                cv_data['education'] = '\n'.join(education_lines[:10])
            if training_lines and not cv_data['training']:
                cv_data['training'] = '\n'.join(training_lines[:15])

            # If no parsed content, use portions of full text
            if not cv_data['employment'] and not cv_data['skills']:
                cv_data['skills'] = full_text[:4000]

            # Escape HTML characters
            for key in cv_data:
                if isinstance(cv_data[key], str):
                    cv_data[key] = escape(cv_data[key])

            # Generate HTML
            html_content = html_template.format(**cv_data)

            # Save HTML file
            html_filename = f'HTML_CVs/{name}.html'
            with open(html_filename, 'w', encoding='utf-8') as f:
                f.write(html_content)

            print(f'Created: {name}.html -> {position}')
            processed += 1

        except Exception as e:
            print(f'Error processing {filename}: {str(e)}')
            import traceback
            traceback.print_exc()

    print(f'\nTotal: {processed} HTML files created in HTML_CVs folder')

if __name__ == '__main__':
    main()
