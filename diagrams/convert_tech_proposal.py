#!/usr/bin/env python3
"""Convert Technical Proposal markdown to Word with images."""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paths
BASE_DIR = "/home/mustafa/Dropbox/Telcobright Customers Work/BTRC/BroadbandPeformanceMonitoringEOI"
MD_FILE = f"{BASE_DIR}/diagrams/Project Understanding Summary-2 (Technical Proposal).md"
OUTPUT_FILE = f"{BASE_DIR}/diagrams/Technical-Proposal.docx"
SCREENSHOTS_DIR = f"{BASE_DIR}/submission/doc/final-zip/specific-exp-cisp/screenshots"

# Image mapping: section keyword -> list of (image file, caption)
BPMN_DIR = f"{BASE_DIR}/diagrams"

IMAGES = {
    "2.1 Architecture Overview": [
        (f"{SCREENSHOTS_DIR}/001.architecture01HighLevel", "Figure: High-Level Architecture"),
        (f"{SCREENSHOTS_DIR}/002.architectureHighLevel-DC_DR", "Figure: DC/DR Architecture"),
    ],
    "2.2 Secure Network Connectivity": [
        (f"{SCREENSHOTS_DIR}/SecureOverlayNetwork", "Figure: Secure IPSec Overlay Network"),
    ],
    "7. Process Flows": [
        (f"{BPMN_DIR}/bpmn-01-onboarding", "BPMN 1: ISP Onboarding Process"),
        (f"{BPMN_DIR}/bpmn-02-data-acquisition", "BPMN 2: Data Acquisition Process"),
        (f"{BPMN_DIR}/bpmn-03-alert-generation", "BPMN 3: QoS Alert Generation Process"),
    ],
}

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_from_md(doc, table_text):
    """Parse markdown table and add to document."""
    lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
    if len(lines) < 2:
        return

    # Parse header and rows
    headers = [c.strip() for c in lines[0].split('|') if c.strip()]
    rows = []
    for line in lines[2:]:  # Skip header and separator
        if '|' in line:
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                rows.append(cells)

    if not headers:
        return

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header.replace('**', '')
        set_cell_shading(cell, '1F2937')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < len(table.columns):
                cell = table.rows[r_idx + 1].cells[c_idx]
                # Clean markdown formatting
                clean_text = cell_text.replace('**', '').replace('`', '')
                cell.text = clean_text
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()

def add_image(doc, image_path, caption=None, width=6.0):
    """Add image with caption."""
    # Try common extensions
    for ext in ['', '.png', '.jpg', '.jpeg', '.PNG', '.JPG']:
        full_path = image_path + ext
        if os.path.exists(full_path):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(full_path, width=Inches(width))

            if caption:
                cap_para = doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(caption)
                cap_run.font.size = Pt(9)
                cap_run.font.italic = True
            return True
    return False

def process_markdown(doc, md_content):
    """Process markdown content and add to document."""
    lines = md_content.split('\n')
    i = 0
    current_section = ""

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # H1 - Title
        if line.startswith('# '):
            title = line[2:].strip()
            para = doc.add_heading(title, level=0)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # H2 - Major sections
        if line.startswith('## '):
            section = line[3:].strip()
            current_section = section
            doc.add_heading(section, level=1)

            # Check for image insertion
            for key, img_list in IMAGES.items():
                if key in section:
                    for img_path, caption in img_list:
                        doc.add_paragraph()
                        if add_image(doc, img_path, caption):
                            print(f"Added image: {caption}")
            i += 1
            continue

        # H3 - Subsections
        if line.startswith('### '):
            subsection = line[4:].strip()
            doc.add_heading(subsection, level=2)

            # Check for image insertion
            for key, img_list in IMAGES.items():
                if key in subsection:
                    for img_path, caption in img_list:
                        doc.add_paragraph()
                        if add_image(doc, img_path, caption):
                            print(f"Added image: {caption}")
            i += 1
            continue

        # H4 - Sub-subsections
        if line.startswith('#### '):
            subsubsection = line[5:].strip()
            doc.add_heading(subsubsection, level=3)
            i += 1
            continue

        # Code blocks
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_para = doc.add_paragraph()
                code_para.style = 'No Spacing'
                for code_line in code_lines:
                    run = code_para.add_run(code_line + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8)
            i += 1
            continue

        # Tables
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, '\n'.join(table_lines))
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            bullet_text = line.strip()[2:]
            # Clean markdown formatting
            bullet_text = re.sub(r'\*\*(.+?)\*\*', r'\1', bullet_text)
            bullet_text = bullet_text.replace('`', '')
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        # Bold text lines (like **Project:** ...)
        if line.strip().startswith('**') and ':**' in line:
            clean_line = line.strip().replace('**', '')
            para = doc.add_paragraph()
            parts = clean_line.split(':', 1)
            if len(parts) == 2:
                run = para.add_run(parts[0] + ':')
                run.bold = True
                para.add_run(' ' + parts[1].strip())
            else:
                para.add_run(clean_line)
            i += 1
            continue

        # Italic notes (*text*)
        if line.strip().startswith('*') and line.strip().endswith('*') and not line.strip().startswith('**'):
            text = line.strip()[1:-1]
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.italic = True
            run.font.size = Pt(9)
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        if text:
            # Clean markdown formatting
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = text.replace('`', '')
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # Remove links
            para = doc.add_paragraph(text)

        i += 1

def add_table_of_contents(doc):
    """Add a Table of Contents field to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Create TOC field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def main():
    print("Converting Technical Proposal to Word...")

    # Read markdown
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # Add title first (extracted from markdown)
    title_para = doc.add_heading("Technical Proposal", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Online Operation and QoS Monitoring System for Fixed Broadband Service Providers")
    run.font.size = Pt(12)
    run.font.italic = True

    doc.add_paragraph()

    # Add Table of Contents heading
    toc_heading = doc.add_heading("Table of Contents", level=1)

    # Add TOC field
    add_table_of_contents(doc)

    # Add note about updating TOC
    toc_note = doc.add_paragraph()
    toc_note.add_run("(Right-click and select 'Update Field' to refresh TOC in Word)").font.size = Pt(8)
    toc_note.runs[0].font.italic = True

    # Page break after TOC
    doc.add_page_break()

    # Process markdown
    process_markdown(doc, md_content)

    # Add DC/DR image after architecture section (if not already added)
    # This is a manual addition for the second architecture diagram

    # Save document
    doc.save(OUTPUT_FILE)
    print(f"Saved to: {OUTPUT_FILE}")

    # List images found
    print("\nImages in screenshots folder:")
    for f in os.listdir(SCREENSHOTS_DIR):
        print(f"  - {f}")

if __name__ == '__main__':
    main()
