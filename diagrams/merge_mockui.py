#!/usr/bin/env python3
"""Merge ISP Portal Screenshots into Technical Proposal as Mock UI section."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE_DIR = "/home/mustafa/Dropbox/Telcobright Customers Work/BTRC/BroadbandPeformanceMonitoringEOI"
TECH_PROPOSAL = f"{BASE_DIR}/diagrams/Technical-Proposal.docx"
ISP_SCREENSHOTS = f"{BASE_DIR}/isp-mon/ISP-Portal-Screenshots.docx"
OUTPUT_FILE = f"{BASE_DIR}/diagrams/Technical-Proposal.docx"

def main():
    print("Merging ISP Portal Screenshots into Technical Proposal...")

    # Open the technical proposal
    doc = Document(TECH_PROPOSAL)

    # Add a page break before the new section
    doc.add_page_break()

    # Add Mock UI section header
    doc.add_heading("Appendix A: ISP Portal Mock UI Screenshots", level=1)

    intro = doc.add_paragraph()
    intro.add_run("The following screenshots demonstrate the proposed ISP Selfcare Portal interface. ").italic = False
    intro.add_run("These mockups illustrate key features including dashboard, device management, BGP routing, security center, alerts, support ticketing, compliance submissions, incident management, and reporting modules.").italic = True

    doc.add_paragraph()

    # Open the screenshots document to extract images
    screenshots_doc = Document(ISP_SCREENSHOTS)

    # Extract and add each image
    img_count = 0
    for rel in screenshots_doc.part.rels.values():
        if "image" in rel.target_ref:
            img_count += 1

    print(f"Found {img_count} images in ISP Portal Screenshots")

    # Extract images to temp folder
    import zipfile
    import shutil

    temp_dir = "/tmp/isp-portal-images"
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)

    with zipfile.ZipFile(ISP_SCREENSHOTS, 'r') as zip_ref:
        for file in zip_ref.namelist():
            if file.startswith('word/media/'):
                zip_ref.extract(file, temp_dir)

    # Get list of extracted images sorted by number
    media_dir = os.path.join(temp_dir, 'word', 'media')
    images = sorted(os.listdir(media_dir), key=lambda x: int(x.replace('image', '').replace('.png', '').replace('.jpeg', '').replace('.jpg', '')))

    print(f"Adding {len(images)} images to Technical Proposal...")

    # Add images with captions
    for i, img_file in enumerate(images, 1):
        img_path = os.path.join(media_dir, img_file)

        # Add image
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(img_path, width=Inches(5.5))

        # Add caption
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(f"Screenshot {i}: ISP Portal Interface")
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True

        # Add spacing
        doc.add_paragraph()

        if i % 10 == 0:
            print(f"  Added {i} images...")

    # Clean up
    shutil.rmtree(temp_dir)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"Saved to: {OUTPUT_FILE}")
    print(f"Total images added: {len(images)}")

if __name__ == '__main__':
    main()
